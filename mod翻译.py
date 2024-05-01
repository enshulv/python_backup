import os
import re
from docx import Document

path = 'E:\Desktop\新建文件夹 (4)'

def 待翻译doc(text_list,dir_doc):
    doc=Document()
    for a in text_list:
        doc.add_paragraph(a)
    doc.save(dir_doc)

def 读取翻译(dir_doc):
    doc = Document(dir_doc)
    列表 = []
    for a in doc.paragraphs:
        列表.append(a.text.replace('【','[').replace('】',']'))
    os.remove(dir_doc)
    return 列表

dir = os.listdir(os.path.join(path,'待翻译'))
for 文件 in dir:
    文件地址 = os.path.join(path,'待翻译',文件)
    with open(文件地址,'r',encoding='UTF-8') as f:
        text_list=f.readlines()  
    有代码行={}
    无代码行=[]
    翻译替换_字典 = {}
    计数 = 0

    for text in text_list:
        #判断是否有注释
        if text[0] == '#':
            continue
        单行= re.search('(?<= ").*(?=")',text)
        #判断是否有内容
        if not 单行:
            continue

        单行_text=单行.group()
        if 单行_text == '':
            continue

        正则 = '\[.*?\]|\$.*?\$|#[a-z] #[a-z].*?#\!.*?#\!|#.*?#!'
        判断=re.search(正则 ,单行_text)
        #判断是否有代码
        if 判断:
            分割 = re.split('('+正则+')',单行_text)
            #如果不是纯代码内容
            if 分割[0] == '' and 分割[-1] == '' and len(分割) <=3:
                continue
            #分割完之后分别换成可翻译形式，之后再替换回来
            列表=[]
            for 分割项 in 分割:
                辨别 = re.search('('+正则+')',分割项)
                if not 辨别:
                    列表.append(分割项)
                    continue
                代号 = '['+hex(计数)+']'
                翻译替换_字典[代号] = 分割项
                列表.append(代号)
                计数+=1
            有代码行[单行_text] = ''.join(列表)
        else:
            无代码行.append(单行_text)

    doc_有代码 = os.path.join(path,f'{文件}-有代码.docx')
    doc_无代码 = os.path.join(path,f'{文件}-无代码.docx')
    待翻译doc(有代码行.values(),doc_有代码)
    待翻译doc(无代码行,doc_无代码)
    input('翻译完点回车')
    有代码_翻译 = 读取翻译(doc_有代码)
    无代码_翻译 = 读取翻译(doc_无代码)

    翻译_f = open(os.path.join(path,文件),'w',encoding='UTF-8')
    for 原文,译文 in zip(有代码行.keys(),有代码_翻译):
        拆分 = re.split('(\[.*?\])',译文)
        cc= list(翻译替换_字典)
        for 拆分项 in 拆分:
            if 拆分项 in list(翻译替换_字典):
                译文 = 译文.replace(拆分项,翻译替换_字典[拆分项])
                有代码行[原文] = 译文
    无代码行_字典 = {}
    for 原文_,译文_ in zip(无代码行,无代码_翻译):
        无代码行_字典[原文_] = 译文_

    for 内容 in text_list:
        有效内容 = re.search('(?<= ").*(?=")',内容)
        if not 有效内容:
            翻译_f.write(内容)
            continue
        是否存在译文 = False
        有效内容_text = 有效内容.group()
        for 原文_有 in 有代码行.keys():
            if 原文_有 == 有效内容_text:
                结果 = 内容.replace(有效内容_text,有代码行[原文_有])
                翻译_f.write(结果)
                是否存在译文 = True
                break

        for 原文_无 in 无代码行:
            if 原文_无 == 有效内容_text:
                结果 = 内容.replace(有效内容_text,无代码行_字典[原文_无])
                翻译_f.write(结果)
                是否存在译文 = True
                break

        if not 是否存在译文:
            翻译_f.write(内容)
    翻译_f.close()
        

        

                
                        
