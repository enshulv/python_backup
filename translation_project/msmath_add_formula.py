from formula_conversion import to_omml, formula_add
import docx

# Open the LaTeX file and read the content, replacing newlines with nothing
with open('E:\\Desktop\\latex.txt', 'r', encoding='utf-8') as file:
    latex_content = file.read().replace('\n', '')

# Create a new Word document
document = docx.Document()

# Add a new paragraph to the document
paragraph = document.add_paragraph('')

# Add a run to the paragraph to hold the formula
run = paragraph.add_run()

# Convert the LaTeX to MathML
omml = to_omml(latex_content)

# Add the formula to the run
formula_add(omml, run)

# Save the document (you need to specify the path where you want to save the document)
document.save(r'path\to\save\document.docx')