import importlib as imp
from pynput import mouse, keyboard
import time
import requests as req
import re as regex
import hashlib as hashing
import pyperclip as clipboard
from lxml import etree as et
import json
from docx import Document
import docx
import copy
from bs4 import BeautifulSoup as soup
import base64
from docx2python import docx2python as docx_to_python
from tqdm import tqdm

def baidu_translate(text, language):
    url = ''
    appid = ''
    salt = ''
    secret_key = ''
    signature = appid + text + salt + secret_key
    sign = hashing.md5(signature.encode('utf-8')).hexdigest()
    data = {
        'q': text,
        'from': 'auto',
        'to': language,
        'appid': appid,
        'salt': salt,
        'sign': sign,
        'Content-Type': '',
    }
    result = json.loads(req.post(url, data=data).text)['trans_result']
    return result

def google_translate(text, language):
    url = ''
    appid = ''
    salt = ''
    secret_key = ''
    signature = appid + text + salt + secret_key
    sign = hashing.md5(signature.encode('utf-8')).hexdigest()
    data = {
        'q': text,
        'from': 'auto',
        'to': language,
        'appid': appid,
        'salt': salt,
        'sign': sign,
        'Content-Type': '',
    }
    result = json.loads(req.post(url, data=data).text)['trans_result']
    return result

def extract_images(path):
    doc = docx_to_python(path)
    image_dict = doc.images
    return image_dict

def formula_recognition_pdf(path, save_path):
    url = ''
    headers = {
        'app_id': '',
        'app_key': ''
    }
    files = [('file', open(path, 'rb'))]
    while True:
        try:
            pdf_id = json.loads(req.post(url, headers=headers, data={}, files=files).text)['pdf_id']
        except Exception as err:
            print(err)
            print('Uploading PDF failed. Retrying in 60 seconds.')
            time.sleep(60)
            continue
        break
    print(pdf_id)
    while True:
        url_progress = url + pdf_id
        try:
            progress = json.loads(req.get(url_progress, headers=headers).text)['percent_done']
        except Exception:
            print('Error getting progress. Retrying in 60 seconds.')
            time.sleep(60)
            continue
        print('Current progress: ' + str(progress))
        if progress == 100:
            break
        time.sleep(60)
    url_download = url + pdf_id + '.docx'
    while True:
        try:
            download = req.get(url_download, headers=headers).content
        except Exception:
            print('Error downloading result. Retrying in 60 seconds.')
            time.sleep(60)
            continue
        break
    with open(save_path, 'wb') as f:
        f.write(download)

def formula_recognition_image(img):
    img_base64 = "data:image/jpg;base64," + base64.b64encode(img).decode()
    headers = {
        'app_id': '',
        'app_key': '',
        'Content-type': '',
    }
    data = {
        'src': img_base64,
    }
    result = json.loads(req.post('', data=json.dumps(data), headers=headers).text)
    return result

def superscript_subscript(paragraph):
    count = 0
    dictionary = {}
    list = []
    for run in paragraph.runs:
        superscript = run.font.superscript
        subscript = run.font.subscript
        footnote = regex.search('w:footnoteReference', run._element.xml)
        if superscript:
            tag = '[' + str(hex(count)) + ']'
            dictionary[(tag, 'superscript')] = run.text
            count += 1
            run.text = tag
        if subscript:
            tag = '[' + str(hex(count)) + ']'
            dictionary[(tag, 'subscript')] = run.text
            count += 1
            run.text = tag
        if footnote:
            tag = '[' + str(hex(count)) + ']'
            dictionary[(tag, 'footnote')] = et.XML(run._element.xml)
            count += 1
            run.text = tag
        list.append(run.text)
        run.text = run.text.replace(run.text, '')
    content = ''.join(list)
    return content, dictionary

def add_run(paragraph, content):
    run = paragraph.add_run(content)
    run.font.superscript = False
    run.font.subscript = False
    run.font.italic = False
    xml = et.XML(run._element.xml)
    paragraph._element.append(xml)
    run.text = run.text.replace(run.text, '')

def translate_preserve_superscript(paragraph, translation, dictionary):
    split = regex.split(r'(\[0[xX][0-9a-fA-F]+\])', translation.replace('[ 0', '[0'))
    for part in split:
        tag = regex.search(r'\[0[xX][0-9a-fA-F]+\]', part)
        if tag:
            for replace_tag, content in dictionary.items():
                if replace_tag[0] == tag.group().lower():
                    if replace_tag[1] == 'superscript':
                        run = paragraph.add_run(content)
                        run.font.superscript = True
                    if replace_tag[1] == 'subscript':
                        run = paragraph.add_run(content)
                        run.font.subscript = True
                    if replace_tag[1] == 'footnote':
                        paragraph._element.append(content)
        else:
            try:
                add_run(paragraph, part)
            except Exception:
                pass

def translate_normal(paragraph, translation):
    runs = paragraph.runs
    for run in runs:
        if run == runs[0]:
            run.text = translation
        else:
            run.text = ''

def detect(translation):
    if translation == '工具书类':
        translation = '参考文献'
    if '哪里' in translation:
        translation = translation.replace('哪里', '其中')
    if '三。' in translation:
        translation = translation.replace('三。', '3')
    if '无花果。' in translation:
        translation = translation.replace('无花果。', '图')
    if '抽象的。' in translation:
        translation = translation.replace('抽象的。', '摘要。')
    if '肛门。' in translation or '肛门！' in translation:
        translation = translation.replace('肛门', 'Anal.')
    if '不是的。' in translation:
        translation = translation.replace('不是的。', 'No.')
    if '内政部：' in translation:
        translation = translation.replace('内政部：', 'DOI:')
    if '答。' in translation:
        translation = translation.replace('答。', 'A。')
    if 'two。' in translation:
        translation = translation.replace('two。', '2.')
    if '兽人' in translation:
        translation = translation.replace('兽人', 'ORCID')
    if translation == 'A B S T R A C T公司':
        translation = '摘要'
    if translation == 'A R T I C L E in F O公司':
        translation='详细信息'
    if translation == '5~6成熟':
        translation = '中等'
    return translation

def translate_stage(tag, merge, language, en_to_cn):
    if tag == 0:
        try:
            time.sleep(0.1)
            translation = google_translate(merge, language)[0]['dst']
        except Exception:
            try:
                time.sleep(0.2)
                translation = baidu_translate(merge, language)[0]['dst']
            except Exception:
                try:
                    time.sleep(0.1)
                    translation = google_translate(merge, language)[0]['dst']
                except Exception:
                    translation = merge
    if tag == 1:
        try:
            time.sleep(0.2)
            translation = google_translate(merge, language)[0]['dst']
        except Exception:
            try:
                time.sleep(1)
                translation = baidu_translate(merge, language)[0]['dst']
            except Exception:
                try:
                    time.sleep(1)
                    translation = baidu_translate(merge, language)[0]['dst']
                except Exception:
                    translation = merge
    if en_to_cn == 1:
        translation = merge + '[换行]' + translation
    translation = detect(translation)
    return translation

def table_translate(tag, doc, language):
    tables = doc.tables
    for table in tqdm(tables):
        try:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if not regex.search('\S', paragraph.text):
                            continue
                        match = regex.search(r'[a-z]+|[\u4e00-\u9fa5]+|[\u0400-\u04FF]+|[\uAC00-\uD7AF]+|[\u0800-\u4e00]+|[\u3040-\u31FF]+', text)
                        if not match:
                            continue
                        content, dictionary = superscript_subscript(paragraph)
                        result = content.replace('\n', '')
                        translation = translate_stage(tag, result, language, 0)
                        translate_preserve_superscript(paragraph, translation, dictionary)
        except Exception:
            print('Table error')
    print('\nTable translation completed')

def translate(name, language, source_path, save_path, tag, skip, table_engine, en_to_cn):
    doc_path = source_path + '//' + name + '.docx'
    doc = Document(doc_path)
    tables = doc.tables
    table_count = len(tables)
    if table_count != 0:
        table_translate(table_engine, doc, language)
    for paragraph in tqdm(doc.paragraphs):
        if not regex.search('\S', paragraph.text):
            continue
        content = paragraph.text
        result = content.replace('\n', '')
        if skip == 1:
            if result == 'References' or result == 'REFERENCES':
                print('\nReferences skipped')
                break
        indent = paragraph.paragraph_format.left_indent
        if indent is not None:
            if indent < 0:
                paragraph.paragraph_format.left_indent = 0
        translation = regex.sub(r'【(.*?)】', r'[\1]', translate_stage(tag, result, language, en_to_cn))
        translate_normal(paragraph, translation)
    translated_doc_path = save_path + '//' + name + '-translated.docx'
    doc.save(translated_doc_path)