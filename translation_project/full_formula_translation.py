import time
import docx
import sys
sys.path.append(r'')
from translation_tools import google_translate, detect, translate_stage, superscript_subscript, translate_preserve_superscript
import re
from lxml import etree
from docx.shared import Pt
from collections import Counter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.support.wait import WebDriverWait
from tqdm import tqdm

def add_translation(paragraph, dictionary, result):
    for run in paragraph.runs:
        if 'w:object' in run._element.xml:
            continue
        run.clear()
    for a in result:
        if a != '' or a != None:
            tag = re.search(r'\[0[xX][0-9a-fA-F]+\]', a)
            if tag:
                runs = paragraph.add_run('')
                formula = dictionary[tag.group().lower()]
                try:
                    if 'w:r' in str(formula):
                        paragraph._element.append(formula)
                    else:
                        runs._element.append(formula)
                except KeyError:
                    paragraph.add_run('-'*30)
            else:
                paragraph.add_run(a)

def error(content, language):
    try:
        result = detect('\n'.join([a['dst'] for a in google_translate(content, language)]))
        return result
    except Exception as err:
        print(err)
        return content

def translate1(paragraph, language):
    list_ = []
    dictionary = {}
    count = 0
    for b in paragraph._element:
        if 'w:r' in str(b):
            if 'w:object' in str(b.xml):
                index = '[' + str(hex(count)) + ']'
                list_.append(index)
                dictionary[index] = b
                count += 1
                continue
            list_.append(b.text)
        if 'oMath' in str(b):
            index = '[' + str(hex(count)) + ']'
            list_.append(index)
            dictionary[index] = b
            count += 1
    merged = ''.join(list_)
    time.sleep(1)
    translation = re.sub(r'【(.*?)】', r'[\1]', error(merged, language))
    if language == 'en':
        split = re.split(r'(\[0[xX][0-9a-fA-F]+\])', translation.replace('[ 0', '[0'))
    else:
        split = re.split(r'(\[0[xX][0-9a-fA-F]+\])', translation)
    add_translation(paragraph, dictionary, split)

def table_translation(open_doc, language):
    tables = open_doc.tables
    for a in tqdm(tables):
        for b in a.rows:
            for c in b.cells:
                for d in c.paragraphs:
                    count = 0
                    match = re.search(r'[a-z]+|[\u4e00-\u9fa5]+|[\u0400-\u04FF]+|[\uAC00-\uD7AF]+', d.text)
                    if not match:
                        continue
                    xml = d._element.xml
                    if 'm:oMath' in xml or 'w:object' in xml:
                        translate1(d, language)
                    else:
                        for e in d.runs:
                            content = e.text
                            if content == '':
                                continue
                            count += 1
                            if count == 1:
                                try:
                                    time.sleep(0.1)
                                    result = google_translate(d.text.replace('\n', ''), language)[0]['dst']
                                except Exception:
                                    result = d.text
                                    print('\n' + result)
                                e.text = e.text.replace(e.text, result)
                                continue
                            e.text = e.text.replace(e.text, '')
    print('\nTable translation completed')

def run_translation(language, name, location, save):
    open_doc = docx.Document(location + '//' + name + '.docx')
    tables = open_doc.tables
    count = len(tables)
    if count != 0:
        try:
            table_translation(open_doc, language)
        except IndexError:
            print('Table error')
    for a in tqdm(open_doc.paragraphs):
        if not re.search('\S', a.text):
            continue
        xml = a._element.xml
        if 'm:oMath' in xml or 'w:object' in xml:
            translate1(a, language)
        else:
            content, dictionary = superscript_subscript(a)
            result = content.replace('\n', '')
            translation = re.sub(r'【(.*?)】', r'[\1]', translate_stage(0, result, language, 0))
            translate_preserve_superscript(a, translation, dictionary)
    open_doc.save(save + '//' + name + '-translated.docx')