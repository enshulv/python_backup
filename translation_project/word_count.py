import docx
import nltk
import os
import zmail
import math
import re
import time
import shutil
import langid
from imbox import Imbox
from lxml import html
import tkinter.messagebox as tk
import tkinter

def word_count(path):
    list_ = []
    remove_extra = nltk.tokenize.RegexpTokenizer(r'\w+|\d+|\n')
    word_count = 0
    if '.docx' in path[-5:]:
        document = docx.Document(path)
        list_.extend(textbox(document))
        table = table(document)
        if len(table) != 0:
            list_.extend(table(document))
        for a in document.paragraphs:
            list_.append(a.text)
    elif '.txt' in path[-5:]:
        with open(path,'r+',encoding='UTF-8') as open_:
            list_.extend(open_.readlines())
    language = langid.classify(''.join(list_))
    for b in list_:
        if b == '':
            continue
        if language[0] == 'zh':
            tkinter.Tk().withdraw()()
            tk.showinfo('Error','The article is not in English, please enter the word count manually')
            word_count = int(input('Word count:'))
            break
        tokens = remove_extra.tokenize(b)
        word_count += int(len(tokens)) 
    return word_count

def textbox(document):
    list_ = []
    for a in document.element.body.iter():
        if a.tag.endswith('textbox'):
            for b in a.iter():
                if b.tag.endswith('main}r'):
                    list_.append(b.text)
    return list_

def table(document):
    list_ = []
    table = document.tables
    for a in table:
        for b in a.rows:
            for c in b.cells:
                list_.append(c.text)
            
    return list_

def run(path,save):
    while True:
        time.sleep(1)
        to_estimate = os.listdir(path)
        if len(to_estimate) != 0: 
            try:
                email = re.findall(r'.*@.*\.com',to_estimate[0].replace('邮','.'))[0]
            except Exception:
                os.remove(path+to_estimate[0])
                continue
            print(email+'，how many copies:')
            copies = input()
            if copies == '\q':
                continue
            price = '0'+str(copies)
            message = 'Your price is: '+price+'\nPlease provide it to the customer service'
            email = {
                'subject':'Estimate',
                'content_text':message,
            }
            try:
                login = zmail.server('','')
                login.send_mail(email,email)
            except Exception as err:
                print(err)
                time.sleep(1)
            to_estimate = os.listdir(path)
            for a in to_estimate:
                if email.replace('.','邮') in a:
                    try:
                        shutil.move(path+a,save)
                    except Exception:
                        input('Press when you are done')
                        os.remove(path+a)

def move_file(original_path,new_path):
    with open(original_path,'rb') as f:
        content = f.read()
    with open(new_path,'wb') as f:
        f.write(content)
    os.remove(original_path)

def send_results(path,move):
    while True:
        to_send = os.listdir(path)
        if len(to_send) !=0:
            classified = classify(to_send)
            try:
                login = zmail.server('','')
            except Exception as err:
                print(err)
                time.sleep(1)
            for a,b in classified.items():
                file_paths = [path+c for c in b]
                move_paths = [move+cc for cc in b]
                print('Sending: '+a)
                email = {
                'subject':'Translation results',
                'content_text':'Please check, if there are any translation or formatting issues, please reply to the email. If it cannot be fixed, you can apply for a full refund',
                'attachments':file_paths,
                }
                try:
                    login.send_mail(a,email)
                except Exception as err:
                    print(err)
                for d,e in zip(file_paths,move_paths):
                    move_file(d,e)
                time.sleep(1)
        time.sleep(60)

def classify(file_directory):
    list_ = []
    for a in file_directory:
        try:
            email = re.findall(r'.*@.*\.com|.*@.*\.cn',a.replace('邮','.'))[0]
            if 'en-' in email:
                email = email.replace('en-','')
            if 'nore-' in email:
                email = email.replace('nore-','')
            if 'qk-' in email:
                email = email.replace('qk-','')
            if 'ew-' in email:
                email = email.replace('ew-','')
        except IndexError:
            continue
        list_.append(email)
    unique = set(list_)
    dict_ = {}
    for b in unique:
        list_2 = []
        for c in file_directory:
            if b in c.replace('邮','.'):
                list_2.append(c)
        dict_[b] = list_2
    return dict_

def send(email,copies,flag):
    try:
        if '163.' in email:
            login = zmail.server('', '')
        else:
            login = zmail.server('','')
    except Exception as err:
        print(err)
        time.sleep(3)
    if flag == 0:
        price_code = '0' + str(copies)
        message = 'Please write the price code in the "Order Remark" when paying: '+email+'\nThe price code is: '+price_code+'\nPlease tell the customer service the price code'
    else:
        message = 'Original layout, buy '+str(copies)+' copies'
    email = {
        'subject':'Estimate',
        'content_text':message,
    }
    try:
        login.send_mail(email,email)
    except Exception as err:
        print(err)
        time.sleep(1)

def check_payment(txt):
    while True:
        with Imbox('','','',ssl=True) as login:
            unread_mails = login.messages(unread=True)
            for uid,mail in unread_mails:
                content = mail.body['plain'][0]
                sender = mail.sent_from[0]['email']
                if 'paid' in content:
                    with open(txt,'r+') as f:
                        read = f.read()
                        if sender not in read:
                            print('Someone has paid - '+sender)
                            f.write(sender+'\n')
                login.mark_seen(uid)
        time.sleep(300)

def transfer(txt,path,paid):
    with open(txt) as f:
        read = f.readlines()
        list_ = []
        for a in read:
            results = os.listdir(path)
            for b in results:
                if a.replace('.','邮').replace('\n','') in b:
                    shutil.move(path+b,paid)

def auto_estimate(path, flag):
    while True:
        directory = os.listdir(path)
        if len(directory) != 0:
            cleaned = classify(directory)
            for a, b in cleaned.items():
                list_ = []
                for c in b:
                    file = path + c
                    try:
                        word_count = word_count(file)
                    except Exception as err:
                        word_count = 0
                    list_.append(word_count)
                    os.remove(file)
                if 0 in list_:
                    print(str(a)+' has a word count of 0, it might be Chinese, please check manually')
                    time.sleep(10)
                    continue
                total_word_count = sum(list_)
                # reduced = total_word_count - 300
                copies = math.ceil(total_word_count/3500)
                send(a, copies, flag)
                time.sleep(10)
        time.sleep(1)

def pick(path):
    with open(r'') as f:
        read = f.read()
        email_ids = read.split('\n')
        for b in email_ids:
            file_names = os.listdir(path)
            for c in file_names:
                replaced = c.replace('.','邮')
                if b in replaced:
                    print(b)

def full_auto_process(path):
    while True:
        with Imbox('','','',ssl=True) as login:
                unread_mails = login.messages(unread=True)
                for uid,mail in unread_mails:
                    attachments = mail.attachments
                    sender = mail.sent_from[0]['email']
                    if len(attachments) != 0:
                        for a in attachments:
                            attachment_name = sender.replace('.','邮')+'-'+a['filename']
                            content = a['content']
                            with open(path+attachment_name,'wb') as f:
                                f.write(content.getvalue())
                    login.mark_seen(uid)
        time.sleep(60)


def word_calculation(files,path):
    sum_ = []
    for a in files:
        exclude = re.findall(r'template|format',a)
        if exclude:
            continue
        if 'html' in a:
            open_ = html.parse(path+'\\'+a)
            parse = open_.xpath('//div [@id="page-container"]')
            list_ = []
            for a in parse:
                list_.append(a.xpath('string()'))
            content = ''.join(list_)
            remove_extra = nltk.tokenize.RegexpTokenizer(r'\w+|\d+')
            sum_.append(int(len(remove_extra.tokenize(content))))
        if 'docx' in a:
            sum_.append(estimate(path+'\\'+a))
    reduced = sum(sum_)-300
    copies = math.ceil(reduced/5000)
    return copies