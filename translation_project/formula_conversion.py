from translation_tools import formula_recognition_image, extract_images
import os
import time
import docx
from lxml import etree
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2python import docx2python as doc2
from docx.shared import Pt
import re
import mathml2omml
import json
import requests as req
import base64
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from collections import Counter

def style(write, content):
    write.paragraph_format.line_spacing = 1.5
    write.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = write.add_run(content)
    run.font.name = '宋体'
    run.font.size = Pt(12)

def merge_content(open_doc):
    images = [[a for a in extract_images(path).keys()], [b for b in extract_images(path).values()]]
    for c in open_doc.paragraphs:
        content = c.text
        xml = c._element.xml
        if content != '':
            write = open_doc.add_paragraph('')
            style(write, content)
        if content == '' and 'w:drawing' in xml:
            write = open_doc.add_paragraph('')
            write.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = write.add_run('')
            with open(images[0][0], 'wb') as aaa:
                aaa.write(images[1][0])
            run.add_picture(images[0][0])
            images[0].pop(0)
            images[1].pop(0)
            continue
    open_doc.save(r'')

def convert_formula(file):
    open_doc = docx.Document()
    for a in extract_images(file).values():
        try:
            latex = '$$\n' + formula_recognition_image(a) + '\n$$'
            print(latex)
            input()
        except KeyError:
            open_doc.add_paragraph('This formula is missing')
            time.sleep(0.5)
            continue
        open_doc.add_paragraph(latex)
        open_doc.save(r'')
        time.sleep(1)

def filter_images(file):
    open_doc = docx.Document()
    for a, b in extract_images(file).items():
        with open(a, 'wb') as aa:
            aa.write(b)
        open_doc.add_picture(a, width=docx.shared.Cm(12))
        time.sleep(0.5)
        os.remove(a)
    open_doc.save(r'')

def background_images(open_doc, file_path):
    dictionary = {}
    image_open = doc2(file_path)
    images = image_open.images
    part = open_doc.part._rels
    for a in open_doc.paragraphs:
        list = []
        for b in a.runs:
            xml = b._element.xml
            filter_img = re.findall(r'wp:simplePos|wp:positionH|wp:positionV', xml)
            id = re.findall(r'(?<=<a:blip r:embed=").*(?="/>)', xml)
            if filter_img and len(id) != 0:
                for img_id, img_content in part.items():
                    if img_id == id[0]:
                        name = img_content.target_ref
                        for c, d in images.items():
                            if c in name:
                                list.append((d, b))
        if len(list) != 0:
            dictionary[a] = list
    return dictionary

def recognize_image(img):
    open = "data:image/jpg;base64," + base64.b64encode(img).decode()
    headers = {
        'app_id': '',
        'app_key': '',
        'Content-type': '',
    }
    dictionary = {
        'src': open,
        'formats': ['text', 'data'],
        'data_options': {'include_mathml': True, 'include_latex': True}
    }
    request = json.loads(req.post('', data=json.dumps(dictionary), headers=headers).text)
    return request

def mathml_group(request):
    mathml = []
    latex = []
    for a in request['data']:
        if a['type'] == 'mathml':
            mathml.append(a['value'])
        if a['type'] == 'latex':
            latex.append(a['value'])
    paired = zip(mathml, latex)
    return paired

def superscript_subscript(omml):
    wave = re.search('<m:lim>(.*?)</m:lim>', omml)
    if wave:
        replace = re.sub('(?<=<m:t>)~', '˜', wave.group())
    return replace

def formula_add(omml, runs):
    # add = runs.add_run('')
    namespace = '<m:name xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
    end = '</m:name>'
    formula = namespace + omml + end
    formula = re.sub(r'<m:groupChr><m:groupChrPr>(.*?)</m:groupChr>(.*?)</m:groupChr>', r'<m:acc><m:accPr>\1</m:accPr>\2</m:acc>', formula)
    node = runs._element
    try:
        parsed = etree.XML(formula)[0]
        node.append(parsed)
    except Exception as err:
        print('\n' + str(err) + 'Be careful of missing characters at the end of a line')
        runs.add_run('[' + '-' * 20 + ']')

def textbox(w, h, x, y, xml):
    box = '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">' \
          '<w:pPr><w:framePr w:w="' + w + '" w:h="' + h + '" ' \
          'w:hRule="exact" w:wrap="around" w:vAnchor="page" ' \
          'w:hAnchor="page" w:x="' + x + '" w:y="' + y + '"/></w:pPr></w:document>'
    parsed = etree.XML(box)[0]
    xml._element.append(parsed)

def to_omml(formula):
    omml = mathml2omml.convert(formula)
    wave_bar = re.findall(r'<m:lim>.*?</m:lim>', omml)
    if len(wave_bar) != 0:
        for bar in wave_bar:
            if '_' in bar:
                omml = omml.replace(bar, bar.replace('_', '¯'))
            if '¯' in bar:
                omml = omml.replace(bar, bar.replace('¯', '͟'))
            if '˙' in bar:
                omml = omml.replace(bar, bar.replace('˙', '.'))
    return omml

def add_formula(paragraph, segments, pairs, font_size):
    tuple_list = []
    for item in pairs:
        tuple_list.append(item)
    for segment in segments:
        if segment != '' and segment is not None:
            inline_formula = re.search(r'\\\(.*?\\\)', segment)
            large_formula = re.search(r'\\\[[\d\D]+?\\\]', segment)
            if large_formula:
                remove_frame = re.sub(r'\\\[|\\\]|\n', '', segment)
                for pair in tuple_list:
                    if pair[1] == remove_frame:
                        formula_add(to_omml(pair[0]), paragraph)
                        tuple_list.remove(pair)
                        break
            elif inline_formula:
                remove_frame = re.sub(r'\\\( | \\\)', '', segment)
                for pair in tuple_list:
                    if pair[1] == remove_frame:
                        try:
                            formula_add(to_omml(pair[0]), paragraph)
                        except Exception as err:
                            paragraph.add_run('-' * 20)
                            print(err)
                        tuple_list.remove(pair)
                        break
            else:
                paragraph.add_run(segment)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(font_size)

def frame(paragraph):
    xml = paragraph._element.xml
    coordinates = re.findall(r'<wp:positionH[\d\D]+</wp:positionH>|<wp:positionV[\d\D]+</wp:positionV>|<wp:extent .*/>', xml)
    for coordinate_xml in coordinates:
        if 'wp:positionH' in coordinate_xml:
            x = str(int(re.findall(r'\d+', coordinate_xml)[0]) / 635)
        if 'wp:positionV' in coordinate_xml:
            y = str(int(re.findall(r'\d+', coordinate_xml)[0]) / 635)
        if 'wp:extent' in coordinate_xml:
            h = str(int(re.findall(r'(?<=cy=")\d+(?=")', coordinate_xml)[0]) / 635)
            w = str(int(re.findall(r'(?<=cx=")\d+(?=")', coordinate_xml)[0]) / 635)
    return w, h, x, y

def replace(paragraph, image, font_size):
    time.sleep(0.3)
    result = recognize_image(image)
    split = re.split(r'(\\\(.*?\\\))|(\\\[[\d\D]+?\\\])', result['text'])
    add_formula(paragraph, split, mathml_group(result), font_size)

def font_size_statistics(open_doc):
    list = []
    for a in open_doc.paragraphs:
        try:
            font_size = a.style.font.size.pt
        except AttributeError:
            continue
        list.append(font_size)
    statistics = Counter(list).most_common()[0][0]
    return statistics

def convert_to_english(path, save_path, table):
    open_doc = docx.Document(path)
    font_size = font_size_statistics(open_doc)
    image_count = 0
    if table == 1:
        single_layout = SingleLayoutFormula(path, save_path, open_doc)
        single_layout.table_processing()
    for paragraph, images in background_images(open_doc, path).items():
        image_count += len(images)
        print('\rCurrent image count: ' + str(image_count), end='')
        if len(images) != 1:
            for a in images:
                w, h, x, y = frame(a[1])
                paragraph = paragraph.insert_paragraph_before()
                textbox(w, h, x, y, paragraph)
                try:
                    replace(paragraph, a[0], font_size)
                except Exception as err:
                    print('\n' + str(err))
                    continue
                a[1].clear()
        else:
            w, h, x, y = frame(paragraph)
            textbox(w, h, x, y, paragraph)
            try:
                replace(paragraph, images[0][0], font_size)
            except Exception as err:
                print('\n' + str(err))
                open_doc.save(save_path)
                continue
            images[0][1].clear()
        open_doc.save(save_path)

class SingleLayoutFormula:
    def __init__(self, path, save_path, open_doc=False):
        self.path = path
        self.save_path = save_path
        if open_doc != False:
            self.open_doc = open_doc
        else:
            self.open_doc = docx.Document(self.path)
        self.images = doc2(self.path).images
        self.part = self.open_doc.part._rels
        self.tables = self.open_doc.tables
        self.font_size = font_size_statistics(self.open_doc)
        self.image_count = 0

    def single_layout_image(self, paragraph):
        for a in paragraph.runs:
            if a.text == '':
                xml = a._element.xml
                id = re.findall(r'(?<=<a:blip r:embed=").*(?="/>)', xml)
                if len(id) != 0:
                    self.image_count += 1
                    print('\rImage count: ' + str(self.image_count), end='')
                    for img_id, img_content in self.part.items():
                        if img_id == id[0]:
                            name = img_content.target_ref
                            for c, d in self.images.items():
                                if c in name:
                                    a.clear()
                                    try:
                                        replace(paragraph, d, self.font_size)
                                    except Exception:
                                        with open(c, 'wb+') as img:
                                            img.write(d)
    def table_open(self):
        list = []
        for a in self.tables:
            for b in a.rows:
                for c in b.cells:
                    for d in c.paragraphs:
                        if d.text == '':
                            list.append(d)
        return list

    def table_processing(self):
        table = self.table_open()
        for table_paragraph in table:
            self.single_layout_image(table_paragraph)
            self.open_doc.save(self.save_path)

    def convert_to_english_non_frame(self):
        if len(self.tables) != 0:
            self.table_processing()

        for a in self.open_doc.paragraphs:
            if not re.search(r'CENTER|None', str(a.alignment)):
                self.single_layout_image(a)
                self.open_doc.save(self.save_path)

def run(judge, path, save_path, table):
    if judge == 0:
        convert_to_english(path, save_path, table)
    else:
        # Single layout, remember to center formulas that do not need to be recognized first
        single_layout = SingleLayoutFormula(path, save_path)
        single_layout.convert_to_english_non_frame()

if __name__ == '__main__':
    name = input('Enter name: ')
    path = r'' + name + '.docx'
    save_path = r'' + name + '.docx'
    # The first parameter is 0 for frames/1 for no frames, the second parameter is 1 for frames to process tables
    run(1, path, save_path, 1)