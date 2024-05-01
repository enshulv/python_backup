from 翻译工具 import 公式识别_图像,提取图片
import os
import time
import docx
from lxml import etree
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2python import docx2python as doc2
from docx.shared import Pt
import re
import mathml2omml
import json
import requests as ree
import base64
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from collections import Counter

def 样式(写入,内容):
    写入.paragraph_format.line_spacing=1.5
    写入.alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run=写入.add_run(内容)
    run.font.name = '宋体'
    run.font.size = Pt(12)

def 合并内容(打开):
    图片=[[a for a in 提取图片(地址).keys()],[b for b in 提取图片(地址).values()]]
    for c in 打开.paragraphs:
        内容=c.text
        xml = c._element.xml
        if 内容 != '':
            写入 = 打开.add_paragraph('')
            样式(写入,内容)
        if 内容 == '' and 'w:drawing' in xml:
            写入 = 打开.add_paragraph('')
            写入.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            run=写入.add_run('')
            with open(图片[0][0],'wb') as aaa:
                aaa.write(图片[1][0])
            run.add_picture(图片[0][0])
            图片[0].pop(0)
            图片[1].pop(0)
            continue
    打开.save(r'')

def 转换公式(文件):
    打开=docx.Document()
    for a in 提取图片(文件).values():
        try:
            latex='$$\n'+公式识别_图像(a)+'\n$$'
            print(latex)
            input()
            #latex=公式识别_图像(a)
        except KeyError:
            打开.add_paragraph('这个公式不见乐乐乐乐乐')
            time.sleep(0.5)
            continue
        打开.add_paragraph(latex)
        打开.save(r'')
        time.sleep(1)

def 筛选图片(文件):
    打开 = docx.Document()
    for a,b in 提取图片(文件).items():
        with open(a,'wb') as aa:
            aa.write(b)
        打开.add_picture(a,width=docx.shared.Cm(12))
        time.sleep(0.5)
        os.remove(a)
    打开.save(r'')

def 背景图片(打开,文件地址):
    字典={}
    图片打开 = doc2(文件地址)
    图片 = 图片打开.images
    part = 打开.part._rels
    for a in 打开.paragraphs:
        列表=[]
        for b in a.runs:
            xml = b._element.xml
            筛图 = re.findall(r'wp:simplePos|wp:positionH|wp:positionV', xml)
            id = re.findall(r'(?<=<a:blip r:embed=").*(?="/>)', xml)
            if 筛图 and len(id) != 0:
                for 图id, 图内容 in part.items():
                    if 图id == id[0]:
                        名字 = 图内容.target_ref
                        for c, d in 图片.items():
                            if c in 名字:
                                列表.append((d,b))
        if len(列表) != 0:
            字典[a]=列表
    return 字典

def 识别图像(img):
    打开="data:image/jpg;base64," + base64.b64encode(img).decode()
    头 = {
            'app_id': '',
            'app_key': '',
            'Content-type':'',
    }
    字典={
            'src':打开,
            'formats': ['text', 'data'],
            'data_options': {'include_mathml':True,'include_latex':True}
            }
    请求=json.loads(ree.post('',data=json.dumps(字典),headers=头).text)
    return 请求

def mathml组(请求):
    mathml=[]
    latex=[]
    for a in 请求['data']:
        if a['type'] == 'mathml':
            mathml.append(a['value'])
        if a['type'] == 'latex':
            latex.append(a['value'])
    并组=zip(mathml,latex)
    return 并组

def 上缀(omml):
    波浪=re.search('<m:lim>(.*?)</m:lim>',omml)
    if 波浪:
        替换=re.sub('(?<=<m:t>)~','˜',波浪.group())
    return 替换

def 公式添加(omml,runs):
    #omml=re.sub(r'/></m:groupChr>','/></m:groupChrPr>',omml)
    #上缀(omml)
    命名空间 = '<m:name xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
    结尾 = '</m:name>'
    公式_=命名空间+omml+结尾
    公式=re.sub(r'<m:groupChr><m:groupChrPr>(.*?)</m:groupChr>(.*?)</m:groupChr>',r'<m:acc><m:accPr>\1</m:accPr>\2</m:acc>',公式_)
    节点 = runs._element
    try:
        解析 = etree.XML(公式)[0]
        节点.append(解析)
    except Exception as err:
        print('\n'+str(err)+'小心一行后面少字')
        runs.add_run('[' + '-' * 20 + ']')

def 文本框(w,h,x,y,xml):
    框='<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">' \
      '<w:pPr><w:framePr w:w="'+w+'" w:h="'+h+'" ' \
      'w:hRule="exact" w:wrap="around" w:vAnchor="page" ' \
      'w:hAnchor="page" w:x="'+x+'" w:y="'+y+'"/></w:pPr></w:document>'
    解析 = etree.XML(框)[0]
    xml._element.append(解析)

def 转omml(公式):
    omml = mathml2omml.convert(公式)
    上标_杠=re.findall(r'<m:lim>.*?</m:lim>',omml)
    if len(上标_杠)!=0:
        for 杠 in 上标_杠:
            if '_' in 杠:
                omml=omml.replace(杠,杠.replace('_','¯'))
            if '¯' in 杠:
                omml=omml.replace(杠,杠.replace('¯','͟'))
            if '˙' in 杠:
                omml=omml.replace(杠,杠.replace('˙','.'))
    return omml

def 添加公式(段落,分隔,并组,字号):
    元祖=[]
    for 迭代 in 并组:
        元祖.append(迭代)
    for a in 分隔:
        if a != '' and a != None:
            段内公式 = re.search(r'\\\(.*?\\\)', a)
            大公式 = re.search(r'\\\[[\d\D]+\\\]', a)
            if 大公式:
                去框 = re.sub(r'\\\[|\\\]|\n', '', a)
                #添加 = 段落.add_run('')
                for b in 元祖:
                    if b[1] == 去框:
                        公式添加(转omml(b[0]), 段落)
                        元祖.remove(b)
                        break
            elif 段内公式:
                去框 = re.sub(r'\\\( | \\\)', '', a)
                #添加 = 段落.add_run('')
                for b in 元祖:
                    if b[1] == 去框:
                        try:
                            公式添加(转omml(b[0]), 段落)
                        except Exception as err:
                            段落.add_run('-'*20)
                            print(err)
                        元祖.remove(b)
                        break
            else:
                段落.add_run(a)
    段落.paragraph_format.line_spacing = 1
    段落.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    段落.style.font.name=u'Times New Roman'
    段落.style.font.size =Pt(字号)

def 加框(段落):
    xml=段落._element.xml
    坐标 = re.findall(r'<wp:positionH[\d\D]+</wp:positionH>|<wp:positionV[\d\D]+</wp:positionV>|<wp:extent .*/>', xml)
    for 坐标xml in 坐标:
        if 'wp:positionH' in 坐标xml:
            x=str(int(re.findall(r'\d+',坐标xml)[0])/635)
        if 'wp:positionV' in 坐标xml:
            y=str(int(re.findall(r'\d+',坐标xml)[0])/635)
        if 'wp:extent' in 坐标xml:
            h=str(int(re.findall(r'(?<=cy=")\d+(?=")',坐标xml)[0])/635)
            w =str(int(re.findall(r'(?<=cx=")\d+(?=")',坐标xml)[0])/635)
    return w,h,x,y

def 替换(段落,图片,字号):
    time.sleep(0.3)
    结果=识别图像(图片)
    分隔 = re.split(r'(\\\(.*?\\\))|(\\\[[\d\D]+?\\\])', 结果['text'])
    添加公式(段落, 分隔,mathml组(结果),字号)

def 字号统计(打开):
    列表=[]
    for a in 打开.paragraphs:
        try:
            字号=a.style.font.size.pt
        except AttributeError:
            continue
        列表.append(字号)
    统计=Counter(列表).most_common()[0][0]
    return 统计

def 英文转换(地址,保存,表格):
    打开=docx.Document(地址)
    字号=字号统计(打开)
    图数=0
    if 表格 == 1:
        单排版 = 单排版公式(地址, 保存,打开)
        单排版.表格_处理()
    for 段落,图片 in 背景图片(打开,地址).items():
        图数 += len(图片)
        print('\r目前的图片数量：' + str(图数), end='')
        if len(图片) != 1:
            for a in 图片:
                w, h, x, y = 加框(a[1])
                段落 = 段落.insert_paragraph_before()
                文本框(w, h, x, y, 段落)
                try:
                    替换(段落, a[0], 字号)
                except Exception as err:
                    print('\n'+str(err))
                    continue
                a[1].clear()
        else:
            w, h, x, y = 加框(段落)
            文本框(w, h, x, y, 段落)
            try:
                替换(段落, 图片[0][0], 字号)
            except Exception as err:
                print('\n' + str(err))
                打开.save(保存)
                continue
            图片[0][1].clear()
        打开.save(保存)

class 单排版公式:
    def __init__(self,地址,保存地址,打开=False):
        self.地址=地址
        self.保存地址=保存地址
        if 打开 != False:
            self.打开 =打开
        else:
            self.打开 = docx.Document(self.地址)
        self.图片 = doc2(self.地址).images
        self.part = self.打开.part._rels
        self.表格=self.打开.tables
        self.字号 = 字号统计(self.打开)
        self.图数=0

    def 单排版图片(self,段落):
        for a in 段落.runs:
            if a.text == '':
                xml = a._element.xml
                id = re.findall(r'(?<=<a:blip r:embed=").*(?="/>)', xml)
                if len(id) != 0:
                    self.图数+=1
                    print('\r图片数量：'+str(self.图数),end='')
                    for 图id, 图内容 in self.part.items():
                        if 图id == id[0]:
                            名字 = 图内容.target_ref
                            for c, d in self.图片.items():
                                if c in 名字:
                                    a.clear()
                                    try:
                                        替换(段落, d, self.字号)
                                    except Exception:
                                        with open(c,'wb+') as img:
                                            img.write(d)
    def 表格_打开(self):
        列表 = []
        for a in self.表格:
            for b in a.rows:
                for c in b.cells:
                    for d in c.paragraphs:
                        if d.text == '':
                            列表.append(d)
        return 列表

    def 表格_处理(self):
        表 = self.表格_打开()
        for 表段 in 表:
            self.单排版图片(表段)
            self.打开.save(self.保存地址)

    def 英文转换_非框(self):
        if len(self.表格) !=0:
            self.表格_处理()

        for a in self.打开.paragraphs:
            if not re.search(r'CENTER|None', str(a.alignment)):
                self.单排版图片(a)
                self.打开.save(self.保存地址)

def run(判断,地址,保存,表格):
    if 判断 == 0:
        英文转换(地址, 保存,表格)
    else:
        # 单排版记得先把不需要识别的公式居中
        单排版 = 单排版公式(地址, 保存)
        单排版.英文转换_非框()

if __name__ == '__main__':
    输入=input('输入名字：')
    地址=r''+输入+'.docx'
    保存=r''+输入+'.docx'
    #第一个=0带框/1不带框，第二个=1带框时处理表格
    run(1,地址,保存,1)

