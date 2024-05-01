import docx
import nltk
import os
import zmail
import math
import re
import time
import shutil
import langid
from imbox import Imbox
from lxml import html
import tkinter.messagebox as tk
import tkinter

def 估价(地址):
    列表 = []
    去除多余 = nltk.tokenize.RegexpTokenizer(r'\w+|\d+|\n')
    词数 = 0
    if '.docx' in 地址[-5:]:
        打开=docx.Document(地址)
        列表.extend(文本框(打开))
        表=表格(打开)
        if len(表) != 0:
            列表.extend(表格(打开))
        for a in 打开.paragraphs:
            列表.append(a.text)
    elif '.txt' in 地址[-5:]:
        with open(地址,'r+',encoding='UTF-8') as 打开:
            列表.extend(打开.readlines())
    语种=langid.classify(''.join(列表))
    for b in 列表:
        if b == '':
            continue
        if 语种[0] == 'zh':
            tkinter.Tk().withdraw()()
            tk.showinfo('错误','文章不是英文,手动输入单词')
            词数=int(input('单词数:'))
            break
        分词=去除多余.tokenize(b)
        #分词=nltk.word_tokenize(b)
        词数+=int(len(分词)) 
    return 词数

def 文本框(打开):
    列表=[]
    for a in 打开.element.body.iter():
        if a.tag.endswith('textbox'):
            for b in a.iter():
                if b.tag.endswith('main}r'):
                    列表.append(b.text)
    return 列表

def 表格(打开):
    列表=[]
    表=打开.tables
    for a in 表:
        for b in a.rows:
            for c in b.cells:
                列表.append(c.text)
            
    return 列表

def 运行(地址,保存):
    while True:
        time.sleep(1)
        待估价=os.listdir(地址)
        if len(待估价) != 0: 
            try:
                邮箱=re.findall(r'.*@.*\.com',待估价[0].replace('邮','.'))[0]
            except Exception:
                os.remove(地址+待估价[0])
                continue
            print(邮箱+'，多少份：')
            份数=input()
            if 份数 == '\q':
                continue
            价格='0'+str(份数)
            发送='您的价格为：'+价格+'\n提供给客服即可\n自动估价，请勿回复'
            邮件={
                'subject':'估价',
                'content_text':发送,
            }
            try:
                登录=zmail.server('','')
                登录.send_mail(邮箱,邮件)
            except Exception as err:
                print(err)
                time.sleep(1)
                continue
            待估价=os.listdir(地址)
            for a in 待估价:
                if 邮箱.replace('.','邮') in a:
                    try:
                        shutil.move(地址+a,保存)
                    except Exception:
                        input('看完了再点')
                        os.remove(地址+a)

def 转移文件(原地址,转移地址):
    with open(原地址,'rb') as a:
        内容=a.read()
    with open(转移地址,'wb') as aa:
        aa.write(内容)
    os.remove(原地址)

def 发送结果(地址,移动):
    while True:
        待发送=os.listdir(地址)
        if len(待发送) !=0:
            分类=归类(待发送)
            try:
                登录=zmail.server('','')
            except Exception as err:
                print(err)
                time.sleep(1)
            for a,b in 分类.items():
                文件地址=[地址+c for c in b]
                移动地址=[移动+cc for cc in b]
                print('发送：'+a)
                邮件={
                'subject':'翻译结果',
                'content_text':'久等请查收，翻译或格式有问题请邮箱回复，或加V售后，如无法修复可以申请全额退款\n满意您给个好评就行，感谢',
                'attachments':文件地址,
                }
                try:
                    登录.send_mail(a,邮件)
                except Exception as err:
                    print(err)
                    time.sleep(1)
                for d,e in zip(文件地址,移动地址):
                    转移文件(d,e)
                time.sleep(1)
        time.sleep(60)

def 归类(文件目录):
    列表=[]
    for a in 文件目录:
        try:
            邮箱=re.findall(r'.*@.*\.com|.*@.*\.cn',a.replace('邮','.'))[0]
            if 'en-' in 邮箱:
                邮箱=邮箱.replace('en-','')
            if 'nore-' in 邮箱:
                邮箱=邮箱.replace('nore-','')
            if 'qk-' in 邮箱:
                邮箱=邮箱.replace('qk-','')
            if 'ew-' in 邮箱:
                邮箱=邮箱.replace('ew-','')
        except IndexError:
            continue
        列表.append(邮箱)
    去重=set(列表)
    字典={}
    for b in 去重:
        列表2=[]
        for c in 文件目录:
            if b in c.replace('邮','.'):
                列表2.append(c)
        字典[b]=列表2
    return 字典

def 发送(邮箱,份数,标志):
    try:
        if '163.' in 邮箱:
            登录 = zmail.server('', '')
        else:
            登录=zmail.server('','')
    except Exception as err:
        print(err)
        time.sleep(3)
    #付款时请在“订单备注”处填写:'+邮箱+'，一定要写哈，对不上号的话就没法处理\n
    if 标志 == 0:
        价格码 = '0' + str(份数)
        发送='付款时请在“订单备注”处填写:'+邮箱+'\n价格码为：'+价格码+'\n将价格码告知客服即可'
    else:
        发送 = '原排版买十块那个商品' + str(份数) +'份就行\n目前只做原排版和标准论文格式(单列宋体小四)，如需标准格式买十块那个'+str(份数+1)+'份即可\n付款时请在“订单备注”处填写:' + 邮箱+'，标准格式备注再加标准二字就行'
    邮件={
        'subject':'估价',
        'content_text':发送,
    }
    try:
        登录.send_mail(邮箱,邮件)
    except Exception as err:
        print(err)
        time.sleep(1)

def 检测支付(txt):
    while True:
        with Imbox('','','',ssl=True) as 登录:
            未读邮件=登录.messages(unread=True)
            for uid,邮件 in 未读邮件:
                内容=邮件.body['plain'][0]
                发件人=邮件.sent_from[0]['email']
                if '已付款' in 内容:
                    with open(txt,'r+') as a:
                        读取=a.read()
                        if 发件人 not in 读取:
                            print('有人付款-'+发件人)
                            a.write(发件人+'\n')
                登录.mark_seen(uid)
        time.sleep(300)

def 转移(txt,地址,已付款):
    with open(txt) as aa:
        读取=aa.readlines()
        列表=[]
        for a in 读取:
            结果=os.listdir(地址)
            for b in 结果:
                if a.replace('.','邮').replace('\n','') in b:
                    shutil.move(地址+b,已付款)

def 自动估价(地址,标志):
    while True:
        目录=os.listdir(地址)
        if len(目录) != 0:
            清洗=归类(目录)
            for a,b in 清洗.items():
                列表=[]
                for c in b:
                    文件=地址+c
                    try:
                        字数=估价(文件)
                    except Exception as err:
                        字数=0
                    列表.append(字数)
                    os.remove(文件)
                if 0 in 列表:
                    print(str(a)+'的字数为0，可能是中文，手动看看吧')
                    time.sleep(10)
                    continue
                总字数=sum(列表)
                #减少=总字数-300
                份数=math.ceil(总字数/3500)
                发送(a,份数,标志)
                time.sleep(10)
        time.sleep(1)

def 挑拣(地址):
    with open(r'') as a:
        读取=a.read()
        邮箱号=读取.split('\n')
        for b in 邮箱号:
            文件名=os.listdir(地址)
            for c in 文件名:
                替换=c.replace('.','邮')
                if b in 替换:
                    print(b)

def 全自动流程(地址):
    while True:
        with Imbox('','','',ssl=True) as 登录:
                未读邮件=登录.messages(unread=True)
                for uid,邮件 in 未读邮件:
                    附件=邮件.attachments
                    发件人=邮件.sent_from[0]['email']
                    if len(附件) != 0:
                        for a in 附件:
                            附件名=发件人.replace('.','邮')+'-'+a['filename']
                            内容=a['content']
                            with open(地址+附件名,'wb') as f:
                                f.write(内容.getvalue())
                    登录.mark_seen(uid)
        time.sleep(60)


def 单词计算(文件,地址):
    求和=[]
    for a in 文件:
        排除=re.findall(r'模板|格式',a)
        if 排除:
            continue
        if 'html' in a:
            打开=html.parse(地址+'\\'+a)
            解析=打开.xpath('//div [@id="page-container"]')
            列表=[]
            for a in 解析:
                列表.append(a.xpath('string()'))
            内容=''.join(列表)
            去除多余= nltk.tokenize.RegexpTokenizer(r'\w+|\d+')
            求和.append(int(len(去除多余.tokenize(内容))))
        if 'docx' in a:
            求和.append(估价(地址+'\\'+a))
    减少=sum(求和)-300
    份数=math.ceil(减少/5000)
    return 份数
