import imp
from pynput import mouse,keyboard
import time
import requests as re
import re as ree
import hashlib as ha
import pyperclip as py
from lxml import etree as et
import json
from docx import Document
import docx
import copy
from bs4 import BeautifulSoup as be
import base64
from docx2python import docx2python as doc2
from tqdm import tqdm


def 百度(文本,语言):
    url = ''
    appid = ''
    salt = ''
    密钥 = ''
    签名 = appid + 文本 + salt + 密钥
    sign = ha.md5(签名.encode('utf-8')).hexdigest()
    data = {
        'q': 文本,
        'from': 'auto',
        'to': 语言,
        'appid': appid,
        'salt': salt,
        'sign': sign,
        'Content-Type': '',
    }
    结果 = json.loads(re.post(url, data=data).text)['trans_result']
    return 结果

def 谷歌(文本,语言):
    url=''
    appid=''
    salt=''
    密钥=''
    签名=appid+文本+salt+密钥
    sign=ha.md5(签名.encode('utf-8')).hexdigest()
    data={
        'q':文本,
        'from': 'auto',
        'to': 语言,
        'appid':appid,
        'salt':salt,
        'sign':sign,
        'Content-Type':'',
    }
    结果=json.loads(re.post(url,data=data).text)['trans_result']
    return 结果

def 提取图片(地址):
    doc=doc2(地址)
    字典=doc.images
    return 字典

def 公式识别_pdf(地址,储存):
    url=''
    字典={}
    头 = {
        'app_id': '',
        'app_key': ''
    }
    打开=[('file',open(地址,'rb'))]
    while True:
        try:
            id=json.loads(re.post(url,headers=头,data=字典,files=打开).text)['pdf_id']
        except Exception as err:
            print(err)
            print('上传PDF失败，稍后重试')
            time.sleep(60)
            continue
        break
    print(id)
    while True:
        url=''+id
        try:
            进度=json.loads(re.get(url, headers=头).text)['percent_done']
        except Exception:
            print('获取进度错误，稍后重试')
            time.sleep(60)
            continue
        print('目前进度为:'+str(进度))
        if 进度 == 100:
            break
        time.sleep(60)
    url2 = '' + id + '.docx'
    while True:
        try:
            下载 = re.get(url2, headers=头).content
        except Exception:
            print('获取结果错误，稍后重试')
            time.sleep(60)
            continue
        break
    with open(储存,'wb') as a:
            a.write(下载)

def 公式识别_图像(img):
    打开="data:image/jpg;base64," + base64.b64encode(img).decode()
    头 = {
            'app_id': '',
            'app_key': '',
            'Content-type':'',
    }
    字典={
            'src':打开,
            }
    请求=json.loads(re.post('',data=json.dumps(字典),headers=头).text)
    return 请求

def 上下标(段落):
    计数=0
    字典={}
    列表=[]
    for b in 段落.runs:
        上标 = b.font.superscript
        下标 = b.font.subscript
        脚注=ree.search('w:footnoteReference',b._element.xml)
        if 上标:
            序号 = '[' + str(hex(计数)) + ']'
            字典[(序号,'上标')]=b.text
            计数+=1
            b.text=序号
            #b.text = b.text.replace(b.text, 序号)
        if 下标:
            序号 = '[' + str(hex(计数)) + ']'
            字典[(序号,'下标')]=b.text
            计数+=1
            b.text=序号
        if 脚注:
            序号 = '[' + str(hex(计数)) + ']'
            字典[(序号,'脚注')]=et.XML(b._element.xml)
            计数+=1
            b.text=序号
        列表.append(b.text)
        b.text = b.text.replace(b.text, '')
    内容=''.join(列表)
    return 内容,字典

def 添加_run(段落,内容):
    run=段落.runs[0]
    run.add_text(内容)
    run.font.superscript = False
    run.font.subscript = False
    run.font.italic = False
    #run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    xml=et.XML(run._element.xml)
    段落._element.append(xml)
    run.text=run.text.replace(run.text, '')

def 翻译_保留上下标(段落,翻译,字典):
    分割 = ree.split(r'(\[0[xX][0-9a-fA-F]+\])', 翻译.replace('[ 0', '[0'))
    for c in 分割:
        标记 = ree.search(r'\[0[xX][0-9a-fA-F]+\]', c)
        if 标记:
            for 替换标记,内容 in 字典.items():
                if 替换标记[0] == 标记.group().lower():
                    if 替换标记[1] == '上标':
                        run=段落.add_run(内容)
                        run.font.superscript=True
                    if 替换标记[1] == '下标':
                        run = 段落.add_run(内容)
                        run.font.subscript = True
                    if 替换标记[1] == '脚注':
                        段落._element.append(内容)
        else:
            try:
                添加_run(段落,c)
            except Exception:
                pass

def 翻译_普通(段落,翻译):
    runs= 段落.runs
    for a in runs:
        if a == runs[0]:
            a.text=翻译
        else:
            a.text=''
        

def 检测(翻译):
    if 翻译 == '工具书类':
        翻译 = '参考文献'
    if '哪里' in 翻译:
        翻译 = 翻译.replace('哪里', '其中')
    if  '三。' in 翻译:
        翻译 = 翻译.replace('三。','3')
    if  '无花果。' in 翻译:
        翻译 = 翻译.replace('无花果。','图')
    if  '抽象的。' in 翻译:
        翻译 = 翻译.replace('抽象的。','摘要。')
    if  '肛门。' in 翻译 or '肛门！' in 翻译:
        翻译 = 翻译.replace('肛门','Anal.')
    if  '不是的。' in 翻译:
        翻译 = 翻译.replace('不是的。','No.')
    if  '内政部：' in 翻译:
        翻译 = 翻译.replace('内政部：','DOI:')
    if  '答。' in 翻译:
        翻译 = 翻译.replace('答。','A。')
    if  'two。' in 翻译:
        翻译 = 翻译.replace('two。','2.')
    if '兽人' in 翻译:
        翻译 = 翻译.replace('兽人', 'ORCID')
    if 翻译 == 'A B S T R A C T公司':
        翻译='摘要'
    if 翻译 == 'A R T I C L E in F O公司':
        翻译='详细信息'
    if 翻译 == '5~6成熟':
        翻译='中等'
    return 翻译

def 翻译阶段(标记,合并,语言,中英):
    if 标记 == 0:
        try:
            time.sleep(0.1)
            翻译 = 谷歌(合并, 语言)[0]['dst']
        except Exception:
            try:
                time.sleep(0.2)
                翻译 = 腾讯免费(合并, 语言)
            except Exception:
                try:
                    time.sleep(0.1)
                    翻译 = 谷歌(合并, 语言)[0]['dst']
                except Exception:
                    翻译=合并
    if 标记 == 1:
        try:
            time.sleep(0.2)
            #翻译1 = 谷歌(合并, 语言)
            #翻译 = 翻译1[0]['dst']
            翻译 = 腾讯免费(合并, 语言)
        except Exception:
            try:
                time.sleep(1)
                翻译 = 百度(合并, 语言)[0]['dst']
                #翻译 = 腾讯免费(合并, 语言)
            except Exception:
                try:
                    time.sleep(1)
                    翻译 = 百度(合并, 语言)[0]['dst']
                except Exception:
                    翻译=合并
    if 中英 == 1:
        翻译=合并+'[换行]'+翻译
        #翻译 = 翻译 + '\n' + 合并
    翻译=检测(翻译)
    return 翻译

def 表格翻译(标记,打开,语言):
    表格=打开.tables
    for a in tqdm(表格):
        try:
            for b in a.rows:
                for c in b.cells:
                    for d in c.paragraphs:
                        段落 = d.text
                        if not ree.search('\S', d.text):
                            continue
                            # 匹配=ree.findall(r'[A-Za-z]+|[\u4e00-\u9fa5]+|[u0800-u4e00]+',文本)
                        匹配 = ree.search(r'[a-z]+|[\u4e00-\u9fa5]+|[\u0400-\u04FF]+|[\uAC00-\uD7AF]+|[\u0800-\u4e00]+|[\u3040-\u31FF]+', 段落)
                        if not 匹配:
                            continue
                        # 内容, 字典 = 上下标(d)
                        内容 = d.text
                        结果 = 内容.replace('\n', '')
                        翻译 = 翻译阶段(标记, 结果, 语言, 0)
                        # 翻译_保留上下标(d, 翻译, 字典)
                        翻译_普通(d,翻译)
        except Exception:
            print('表格错误')
    print('\n表格翻译完成')

def 翻译(名字,语言,位置,保存地,标记,跳过,表格引擎,中英):
    文档=位置 + '//'+名字 + '.docx'
    打开=docx.Document(文档)
    表格=打开.tables
    数量=len(表格)
    if 数量 != 0:
        表格翻译(表格引擎,打开,语言)
    for a in tqdm(打开.paragraphs):
        if not ree.search('\S', a.text):
            continue
        # 内容, 字典 = 上下标(a)
        内容=a.text
        '''判断=ree.search('[\u4e00-\u9fa5]+',段落1)
        if 判断:
            continue'''
        结果=内容.replace('\n','')
        if 跳过 == 1:
            if 结果 == 'References' or 结果 == 'REFERENCES':
                print('\n已跳过参考文献')
                break
        缩进 = a.paragraph_format.left_indent
        if 缩进 != None:
            if 缩进 < 0:
                a.paragraph_format.left_indent = 0
        翻译=ree.sub(r'【(.*?)】',r'[\1]',翻译阶段(标记,结果,语言,中英))
        # 翻译_保留上下标(a,翻译,字典)
        翻译_普通(a,翻译)
    文档2=保存地+'//'+名字+'-翻译.docx'
    打开.save(文档2)

