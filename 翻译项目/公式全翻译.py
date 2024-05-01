import time
import docx
import sys
sys.path.append(r'C:\Users\25430\Desktop\翻译模块')
from 翻译工具 import 腾讯免费,谷歌,检测,翻译阶段,上下标,翻译_保留上下标
import re
from lxml import etree
from docx.shared import Pt
from collections import Counter
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.support.wait import WebDriverWait
from tqdm import tqdm

def 添加翻译(段落,字典,结果):
    for run in 段落.runs:
        if 'w:object' in run._element.xml:
            continue
        '''if run.text == '\t':
            continue'''
        run.clear()
    #runs = 段落.add_run('\n')
    for a in 结果:
        if a != '' or a != None:
            标记=re.search(r'\[0[xX][0-9a-fA-F]+\]',a)
            if 标记:
                runs=段落.add_run('')
                公式 = 字典[标记.group().lower()]
                try:
                    if 'w:r' in str(公式):
                        段落._element.append(公式)
                    else:
                        runs._element.append(公式)
                except KeyError:
                    段落.add_run('-'*30)
            else:
                段落.add_run(a)
    #段落.style.font.name = u'宋体'

def 报错(内容,语言):
    try:
        #百度翻译
        结果=检测('\n'.join([a['dst'] for a in 谷歌(内容, 语言)]))
        return 结果
    except Exception as err:
        print(err)
        return 内容

def 翻译1(段落,语言):
    列表 = []
    字典 = {}
    计数 = 0
    for b in 段落._element:
        if 'w:r' in str(b):
            if 'w:object' in str(b.xml):
                序号 = '[' + str(hex(计数)) + ']'
                列表.append(序号)
                字典[序号] = b
                计数 += 1
                continue
            列表.append(b.text)
        if 'oMath' in str(b):
            序号 = '[' + str(hex(计数)) + ']'
            列表.append(序号)
            字典[序号] = b
            计数 += 1
    合并 = ''.join(列表)
    time.sleep(1)
    翻译 = re.sub(r'【(.*?)】',r'[\1]',报错(合并,语言))
    if 语言 == 'en':
        分割 = re.split(r'(\[0[xX][0-9a-fA-F]+\])', 翻译.replace('[ 0', '[0'))
    else:
        分割 = re.split(r'(\[0[xX][0-9a-fA-F]+\])', 翻译)
    添加翻译(段落, 字典, 分割)

def 表格翻译(打开,语言):
    表格=打开.tables
    for a in tqdm(表格):
        for b in a.rows:
            for c in b.cells:
                for d in c.paragraphs:
                    计数 = 0
                    匹配 = re.search(r'[a-z]+|[\u4e00-\u9fa5]+|[\u0400-\u04FF]+|[\uAC00-\uD7AF]+', d.text)
                    if not 匹配:
                        continue
                    xml=d._element.xml
                    if 'm:oMath' in xml or 'w:object' in xml:
                        翻译1(d, 语言)
                    #if 'w:object' in xml:
                      #  翻译(d)
                    else:
                        for e in d.runs:
                            内容 = e.text
                            if 内容 == '':
                                continue
                            计数 += 1
                            if 计数 == 1:
                                try:
                                    time.sleep(0.1)
                                    #结果 = 腾讯免费(d.text.replace('\n',''),语言)
                                    结果=谷歌(d.text.replace('\n',' '),语言)[0]['dst']
                                except Exception:
                                    结果=d.text
                                    print('\n'+结果)
                                e.text = e.text.replace(e.text, 结果)
                                continue
                            e.text = e.text.replace(e.text, '')
    print('\n表格翻译完成')

def 运行翻译(语言,名字,位置,保存):
    打开=docx.Document(位置 + '//'+名字 + '.docx')
    表格=打开.tables
    数量=len(表格)
    if 数量 != 0:
        try:
            表格翻译(打开,语言)
        except IndexError:
            print('表格错误')
    for a in tqdm(打开.paragraphs):
        if not re.search('\S', a.text):
            continue
        xml = a._element.xml
        if 'm:oMath' in xml or 'w:object' in xml:
            翻译1(a,语言)
        else:
            内容, 字典 = 上下标(a)
            结果=内容.replace('\n','')
            翻译 = re.sub(r'【(.*?)】',r'[\1]',翻译阶段(0, 结果,语言,0))
            翻译_保留上下标(a, 翻译, 字典)
    打开.save(保存+'//'+名字+'-翻译.docx')
