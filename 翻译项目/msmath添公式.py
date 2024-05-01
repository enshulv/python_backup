from 公式转换 import 转omml,公式添加
import docx

with open('E:\Desktop\latex.txt','r',encoding='utf-8') as a:
    b = a.read().replace('\n','')
doc=docx.Document()
p=doc.add_paragraph('')
run = p.add_run()
omml=转omml(b)
公式添加(omml,run)
doc.save(r'')
