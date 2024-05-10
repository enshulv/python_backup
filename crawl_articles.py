from docx2pdf import convert
from PIL import Image
import requests as req
import re
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.text import WD_BREAK
from latex2mathml import converter as con
import mathml2omml
import os
import time
import tempfile as te
from tqdm import tqdm
import json


class FetchArticle:
    def __init__(self, url: str, storage_dir: str):
        self.docx = Document()
        self.url = url
        self.title_size = 20
        self.author_size = 15
        self.text_size = 12
        self.subtitle_sizes = [6, 5, 4, 3, 2, 1]
        self.storage_dir = storage_dir
        self.title = None
        self.latex_regex = r'(?<=data-tex=").*?(?=")'
        self.docx.styles["Normal"].font.name = "Times New Roman"
        self.docx.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def change_default_font(
        self, english: str = "Times New Roman", chinese: str = "宋体"
    ):
        self.docx.styles["Normal"].font.name = "%s" % english
        self.docx.styles["Normal"]._element.rPr.rFonts.set(
            qn("w:eastAsia"), "%s" % chinese
        )

    def change_text_font(self, text_object):
        english = self.docx.styles["Normal"].font.name
        chinese = self.docx.styles["Normal"]._element.rPr.rFonts.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
        )
        text_object.font.name = "%s" % english
        text_object._element.rPr.rFonts.set(qn("w:eastAsia"), "%s" % chinese)

    def write_title_and_author(self, title, author):
        title_pr = self.docx.add_paragraph("")
        title_pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_pr.add_run(title)
        title_run.font.size = Pt(self.title_size)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        author_pr = self.docx.add_paragraph()
        author_pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_pr.add_run(author)
        author_run.font.size = Pt(self.author_size)
        author_run.bold = True

    def write_text(self, text):
        for a in tqdm(text):
            paragraph = self.docx.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(0)
            original_html = etree.tostring(a, encoding="utf-8").decode("utf-8")
            stripped_tags = re.sub(r"^<.*?>|</{0}>".format(a.tag), "", original_html)
            split_content = list(
                filter(
                    None,
                    re.split(
                        r"(<br/>)|(<a.*?>.*?</a>)|(<.*?>?.*?</.*?>)", stripped_tags
                    ),
                )
            )
            for content in split_content:
                if a.tag == "p" or a.tag == "blockquote":
                    added = self.process_text(paragraph, content)
                    if added:
                        added.font.size = Pt(self.text_size)
                        if a.tag == "blockquote":
                            pPr = OxmlElement("w:pPr")
                            pBdr = OxmlElement("w:pBdr")
                            left = OxmlElement("w:left")
                            left.set(qn("w:val"), "single")
                            left.set(qn("w:color"), "D3D3D3")
                            left.set(qn("w:sz"), "18")
                            left.set(qn("w:space"), "7")
                            pBdr.append(left)
                            pPr.append(pBdr)
                            paragraph._p.append(pPr)
                            paragraph.paragraph_format.space_after = Pt(5)
                            paragraph.paragraph_format.space_before = Pt(5)
                            paragraph.paragraph_format.left_indent = Cm(0.4)
                            added.font.color.rgb = RGBColor(89, 89, 89)
                elif "h" in a.tag:
                    added = self.process_text(paragraph, content)
                    if added:
                        self.change_text_font(added)
                        added.font.size = Pt(
                            self.text_size + self.subtitle_sizes[int(a.tag[-1])]
                        )
                        added.bold = True
            if a.tag == "figure":
                image = req.get(a.find("img").get("data-actualsrc")).content
                with te.TemporaryFile() as img:
                    img.write(image)
                    img_width, img_height = Image.open(img).size
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    page_height = self.docx.sections[0].page_height
                    page_width = self.docx.sections[0].page_width
                    right_margin = self.docx.sections[0].right_margin
                    left_margin = self.docx.sections[0].left_margin
                    top_margin = self.docx.sections[0].top_margin
                    bottom_margin = self.docx.sections[0].bottom_margin
                    width = page_width - (left_margin + right_margin)
                    height = page_height - (top_margin + bottom_margin)
                    img_pr = paragraph.add_run()
                    img_pr.font.size = Pt(self.text_size)
                    if img_width > width / 12700:
                        img_pr.add_picture(img, width=width)
                    elif img_height > height / 12700:
                        img_pr.add_picture(img, height=width)
                    elif img_width > width / 12700 and img_height > height / 12700:
                        img_pr.add_picture(img, height=width, width=width)
                    else:
                        img_pr.add_picture(img)
                img_text = a.find("figcaption")
                if img_text != None:
                    paragraph.add_run(
                        "\n%s" % img_text.text.replace("&gt;", ">").replace("&lt;", "<")
                    )
                time.sleep(0.2)
            if a.tag == "ul":
                for li in a:
                    link = li.find("a")
                    new_pr = self.docx.add_paragraph()
                    new_pr.paragraph_format.line_spacing = 1.5
                    new_pr.paragraph_format.space_after = Pt(0)
                    if link != None:
                        self.add_hyperlink(new_pr, link.text, link.get("href"))
                    else:
                        new_pr.add_run("%s\n" % link.text)
                    self.tab_at_origin(new_pr)

    def process_text(self, paragraph_obj, content):
        added = False
        content.replace("&amp;", "&").replace("&#9;", "\t")
        if "</b>" in content:
            added = paragraph_obj.add_run(re.sub(r"<b>|</b>", "", content))
            added.bold = True
        elif "</i>" in content:
            added = paragraph_obj.add_run(re.sub(r"<i>|</i>", "", content))
            added.italic = True
        elif "<br/>" in content:
            added = paragraph_obj.add_run("\n")
        elif "</a>" in content:
            url = re.findall(r'(?<=href=").*?(?=")', content)[0]
            if "span" in content:
                span_filter = re.findall(r"<span.*?>.*?</span>", content)
                list = []
                for a in span_filter:
                    list.append(re.findall(r"(?<=>).*?(?=</span>)", a)[0])
                text = "".join(list)
            else:
                text = re.findall(r"(?<=>).*?(?=</a>)", content)[0]
            self.add_hyperlink(paragraph_obj, text, url)
        elif "ztext-math" in content:
            added = paragraph_obj.add_run()
            self.add_formula(content, added, self.latex_regex)
        else:
            added = paragraph_obj.add_run(content)
        return added

    def tab_at_origin(self, paragraph_obj):
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numid = OxmlElement("w:numId")
        numid.set(qn("w:val"), "1")
        numPr.append(ilvl)
        numPr.append(numid)
        pPr.append(numPr)
        paragraph_obj._p.append(pPr)

    def add_hyperlink(self, paragraph_obj, text, url):
        id = paragraph_obj.part.relate_to(
            url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), id)
        run = paragraph_obj.add_run(text)
        run.font.size = Pt(self.text_size)
        hyperlink.append(run._element)
        paragraph_obj._p.append(hyperlink)
        run.font.underline = True

    def add_formula(self, formula_dom, runs, latex_regex):
        lat = re.findall(r"{0}".format(latex_regex), formula_dom)[0]
        latex = (
            lat.replace("&gt;", ">")
            .replace("&lt;", "<")
            .replace("&amp;", "&")
            .replace("&#9;", "\t")
        )
        mtl = con.convert(latex)
        try:
            omml = mathml2omml.convert(mtl)
        except Exception:
            print(f"Error, latex is: {latex}")
            mtl = input("Input error content:")
            omml = mathml2omml.convert(mtl)
        namespace = '<m:name xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        end = "</m:name>"
        formula_ = namespace + omml + end
        formula = re.sub(
            r"<m:groupChr><m:groupChrPr>(.*?)</m:groupChr>(.*?)</m:groupChr>",
            r"<m:acc><m:accPr>\1</m:accPr>\2</m:acc>",
            formula_,
        )
        box_format = re.sub(
            r"(<m:rPr>.*?</m:rPr>)",
            r'\1<w:rPr><w:sz w:val="{0}"/><w:szCs w:val="{0}"/></w:rPr>'.format(
                self.text_size * 2
            ),
            formula,
        )
        result = re.sub(
            r"(<m:oMath><m:box>)(.*?)(</m:box></m:oMath>)",
            r'\1<m:boxPr><m:ctrlPr><w:rPr><w:sz w:val="{0}"/><w:szCs w:val="{0}"/></w:rPr></m:ctrlPr></m:boxPr>\2\3'.format(
                self.text_size * 2
            ),
            box_format,
        )
        node = runs._element
        parse = etree.XML(result)[0]
        node.append(parse)

    def write(self, title_xpath, author_xpath, text_xpath):
        if len(self.url) > 1:
            file_naming = input("Save file naming (no need to add suffix):")
        for a in self.url:
            html = etree.HTML(req.get(a).text)
            title = html.xpath(title_xpath)[0]
            author = html.xpath(author_xpath)[0]
            text = html.xpath(text_xpath)
            print("Fetching [%s]" % title)
            self.write_title_and_author(title, author)
            self.write_text(text)
            if len(self.url) > 1:
                runs = self.docx.add_paragraph().add_run()
                runs.add_break(WD_BREAK.PAGE)
                address = os.path.join(self.storage_dir, file_naming) + ".docx"
                self.docx.save(address)
            else:
                address = os.path.join(self.storage_dir, title) + ".docx"
                self.docx.save(address)
            time.sleep(1)


def main():
    # column = req.get('').text
    url = [""]
    # for a in json.loads(column)['data']:
    #     url.append(a['url'])
    dir = r""
    title = '//*[@class="Post-Main Post-NormalMain"]/header/h1/text()'
    author = '//*[@class="UserLink-link"]/text()'
    text = '//div[@options="[object Object]"]/node()'
    fetch = FetchArticle(url, dir)
    fetch.write(title, author, text)


if __name__ == "__main__":
    main()
