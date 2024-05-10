import os
import re
from docx import Document

path = 'E:\Desktop\New Folder (4)'

def translate_doc(text_list, dir_doc):
    doc = Document()
    for text in text_list:
        doc.add_paragraph(text)
    doc.save(dir_doc)

def read_translation(dir_doc):
    doc = Document(dir_doc)
    text_list = []
    for paragraph in doc.paragraphs:
        text_list.append(paragraph.text.replace('【','[').replace('】',']'))
    os.remove(dir_doc)
    return text_list

directory = os.listdir(os.path.join(path, 'Translate'))
for file in directory:
    file_path = os.path.join(path, 'Translate', file)
    with open(file_path, 'r', encoding='UTF-8') as f:
        text_list = f.readlines()
    lines_with_code = {}
    lines_without_code = []
    translation_replace_dict = {}
    count = 0

    for text in text_list:
        # Check for comments
        if text[0] == '#':
            continue
        single_line = re.search('(?<= ").*(?=")', text)
        # Check for content
        if not single_line:
            continue

        single_line_text = single_line.group()
        if single_line_text == '':
            continue

        regex = '\[.*?\]|\$.*?\$|#[a-z] #[a-z].*?#\!.*?#\!|#.*?#!'
        check = re.search(regex, single_line_text)
        # Check for code
        if check:
            split = re.split('(' + regex + ')', single_line_text)
            # If it's not pure code content
            if split[0] == '' and split[-1] == '' and len(split) <= 3:
                continue
            # Split and replace with translatable forms, then replace back
            list = []
            for split_item in split:
                identify = re.search('(' + regex + ')', split_item)
                if not identify:
                    list.append(split_item)
                    continue
                code = '[' + hex(count) + ']'
                translation_replace_dict[code] = split_item
                list.append(code)
                count += 1
            lines_with_code[single_line_text] = ''.join(list)
        else:
            lines_without_code.append(single_line_text)

    doc_with_code = os.path.join(path, f'{file}-with_code.docx')
    doc_without_code = os.path.join(path, f'{file}-without_code.docx')
    translate_doc(lines_with_code.values(), doc_with_code)
    translate_doc(lines_without_code, doc_without_code)
    input('Press Enter when translation is done')
    translation_with_code = read_translation(doc_with_code)
    translation_without_code = read_translation(doc_without_code)

    translation_file = open(os.path.join(path, file), 'w', encoding='UTF-8')
    for original, translation in zip(lines_with_code.keys(), translation_with_code):
        split = re.split('(\[.*?\])', translation)
        for split_item in split:
            if split_item in list(translation_replace_dict):
                translation = translation.replace(split_item, translation_replace_dict[split_item])
                lines_with_code[original] = translation
    lines_without_code_dict = {}
    for original_without, translation_without in zip(lines_without_code, translation_without_code):
        lines_without_code_dict[original_without] = translation_without

    for content in text_list:
        valid_content = re.search('(?<= ").*(?=")', content)
        if not valid_content:
            translation_file.write(content)
            continue
        translation_exists = False
        valid_content_text = valid_content.group()
        for original_with in lines_with_code.keys():
            if original_with == valid_content_text:
                result = content.replace(valid_content_text, lines_with_code[original_with])
                translation_file.write(result)
                translation_exists = True
                break

        for original_without in lines_without_code:
            if original_without == valid_content_text:
                result = content.replace(valid_content_text, lines_without_code_dict[original_without])
                translation_file.write(result)
                translation_exists = True
                break

        if not translation_exists:
            translation_file.write(content)
    translation_file.close()