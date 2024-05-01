from docx2pdf import convert
from PIL import Image
import requests as req
import re
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt,Cm
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement,qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.text import WD_BREAK
from latex2mathml import converter as con
import mathml2omml
import os
import time
import tempfile as te
from tqdm import tqdm
import json
    
class 抓取文章:
    def __init__(self,url:str,储存目录dir:str):
        self.docx = Document()
        self.url = url
        self.标题sz = 20
        self.作者sz = 15
        self.正文sz = 12
        self.小标题sz = [6,5,4,3,2,1]
        self.储存目录dir = 储存目录dir
        self.标题 = None
        self.latex正则 = r'(?<=data-tex=").*?(?=")'
        # self.公式tag = 'span'
        self.docx.styles['Normal'].font.name = 'Times New Roman'
        self.docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    def 修改默认字体(self,英文:str = 'Times New Roman',中文:str='宋体'):
        self.docx.styles['Normal'].font.name = '%s' % 英文
        self.docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'%s' % 中文)

    def 修改文本字体(self,修改对象:object):
        英文 = self.docx.styles['Normal'].font.name
        中文 = self.docx.styles['Normal']._element.rPr.rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
        修改对象.font.name = '%s' % 英文
        修改对象._element.rPr.rFonts.set(qn('w:eastAsia'), u'%s' % 中文)

    def 写入标题和作者(self,标题,作者):
        标题pr=self.docx.add_paragraph('')
        标题pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        标题run=标题pr.add_run(标题)
        标题run.font.size = Pt(self.标题sz)
        标题run.bold = True
        标题run.font.color.rgb = RGBColor(0,0,0)
        作者pr=self.docx.add_paragraph()
        作者pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        作者run=作者pr.add_run(作者)
        作者run.font.size = Pt(self.作者sz)
        作者run.bold = True

    def 写入正文(self,正文):
        for a in tqdm(正文):
            段落 = self.docx.add_paragraph()
            段落.paragraph_format.line_spacing = 1.5
            段落.paragraph_format.space_after = Pt(0)
            原html = etree.tostring(a,encoding = 'utf-8').decode('utf-8')
            去除首尾标签=re.sub(r'^<.*?>|</{0}>'.format(a.tag),'',原html)
            分割内容格式 = list(filter(None,re.split(r'(<br/>)|(<a.*?>.*?</a>)|(<.*?>?.*?</.*?>)',去除首尾标签)))
            for 内容 in 分割内容格式:
                if a.tag == 'p' or a.tag == 'blockquote':
                    添加 = self.处理_正文(段落,内容)
                    if 添加:
                        添加.font.size = Pt(self.正文sz)
                        if a.tag == 'blockquote':
                            pPr=OxmlElement('w:pPr')
                            pBdr = OxmlElement('w:pBdr')
                            left=OxmlElement('w:left')
                            left.set(qn('w:val'),'single')
                            left.set(qn('w:color'),'D3D3D3')
                            left.set(qn('w:sz'),'18')
                            left.set(qn('w:space'),'7')
                            pBdr.append(left)
                            pPr.append(pBdr)
                            段落._p.append(pPr)
                            段落.paragraph_format.space_after = Pt(5)
                            段落.paragraph_format.space_before = Pt(5)
                            段落.paragraph_format.left_indent = Cm(0.4)
                            添加.font.color.rgb = RGBColor(89,89,89)
                elif 'h' in a.tag:
                    添加 = self.处理_正文(段落,内容)
                    if 添加:
                        self.修改文本字体(添加)
                        添加.font.size = Pt(self.正文sz + self.小标题sz[int(a.tag[-1])])
                        添加.bold = True
            if a.tag == 'figure':
                图片 = req.get(a.find('img').get('data-actualsrc')).content
                with te.TemporaryFile() as img:
                    img.write(图片)
                    img_宽,img_长=Image.open(img).size
                    段落.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    页面长度 = self.docx.sections[0].page_height
                    页面宽度 = self.docx.sections[0].page_width
                    右边距 = self.docx.sections[0].right_margin
                    左边距 = self.docx.sections[0].left_margin
                    上边距 = self.docx.sections[0].top_margin
                    下边距 = self.docx.sections[0].bottom_margin
                    宽度 = 页面宽度-(左边距+右边距)
                    长度 = 页面长度-(上边距+下边距)
                    img_pr = 段落.add_run()
                    img_pr.font.size = Pt(self.正文sz)
                    if img_宽 > 宽度/12700:
                        img_pr.add_picture(img,width = 宽度)
                    elif img_长 > 长度/12700:
                        img_pr.add_picture(img,height = 宽度)
                    elif img_宽 > 宽度/12700 and img_长 > 长度/12700:
                        img_pr.add_picture(img,height = 宽度,width = 宽度)
                    else:
                        img_pr.add_picture(img)
                img_text=a.find('figcaption')
                if img_text != None:
                   段落.add_run('\n%s' % img_text.text.replace('&gt;','>').replace('&lt;','<'))
                time.sleep(0.2)
            if a.tag == 'ul':
                for li in a:
                    链接 = li.find('a')
                    new_pr = self.docx.add_paragraph()
                    new_pr.paragraph_format.line_spacing = 1.5
                    new_pr.paragraph_format.space_after = Pt(0)
                    if 链接 != None:
                        self.add_超链接(new_pr,链接.text,链接.get('href'))
                    else:
                        new_pr.add_run('%s\n' % 链接.text)
                    self.制表符_原点(new_pr)

    def 处理_正文(self,段落obj:object,内容:str):
        添加 = False
        内容 = 内容.replace('&gt;','>').replace('&lt;','<').replace('&amp;','&').replace('&#9;','\t')
        if '</b>' in 内容:
            添加 = 段落obj.add_run(re.sub(r'<b>|</b>','',内容))
            添加.bold = True
        elif '</i>' in 内容:
            添加 = 段落obj.add_run(re.sub(r'<i>|</i>','',内容))
            添加.italic = True
        elif '<br/>' in 内容:
            添加 = 段落obj.add_run('\n')  
        elif '</a>' in 内容:
            url=re.findall(r'(?<=href=").*?(?=")',内容)[0]
            if 'span' in 内容:
                span筛 = re.findall(r'<span.*?>.*?</span>',内容)
                列表 = []
                for a in span筛:
                    列表.append(re.findall(r'(?<=>).*?(?=</span>)',a)[0])
                text = ''.join(列表)
            else:
                text=re.findall(r'(?<=>).*?(?=</a>)',内容)[0]
            self.add_超链接(段落obj,text,url) 
        elif 'ztext-math' in 内容:
            添加 = 段落obj.add_run()
            self.添加公式(内容,添加,self.latex正则)
        else:
            添加 = 段落obj.add_run(内容)
        return 添加
    
    def 制表符_原点(self,段落obj:object):
        pPr = OxmlElement('w:pPr')
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'),'0')
        numid = OxmlElement('w:numId')
        numid.set(qn('w:val'),'1')
        numPr.append(ilvl)
        numPr.append(numid)
        pPr.append(numPr)
        段落obj._p.append(pPr)

    def add_超链接(self,段落obj:object, text:str, url:str):
        id=段落obj.part.relate_to(url,RELATIONSHIP_TYPE.HYPERLINK,is_external=True )
        超链接 = OxmlElement('w:hyperlink')
        超链接.set(qn('r:id'),id)
        run = 段落obj.add_run(text)
        run.font.size = Pt(self.正文sz)
        超链接.append(run._element)  
        段落obj._p.append(超链接)
        run.font.underline = True

    def 添加公式(self,公式dom:str,runs:object,latex正则:str):
        lat = re.findall(r'{0}'.format(latex正则),公式dom)[0]
        latex = lat.replace('&gt;','>').replace('&lt;','<').replace('&amp;','&').replace('&#9;','\t')
        mtl = con.convert(latex)
        try:
            omml = mathml2omml.convert(mtl)
        except Exception:
            print(f'出错，错误latex为：{latex}')
            mtl = input('输入错误内容:')
            omml = mathml2omml.convert(mtl)
        命名空间 = '<m:name xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        结尾 = '</m:name>'
        公式_=命名空间+omml+结尾
        公式=re.sub(r'<m:groupChr><m:groupChrPr>(.*?)</m:groupChr>(.*?)</m:groupChr>',r'<m:acc><m:accPr>\1</m:accPr>\2</m:acc>',公式_)
        box格式=re.sub(r'(<m:rPr>.*?</m:rPr>)',r'\1<w:rPr><w:sz w:val="{0}"/><w:szCs w:val="{0}"/></w:rPr>'.format(self.正文sz*2),公式)
        结果=re.sub(r'(<m:oMath><m:box>)(.*?)(</m:box></m:oMath>)',r'\1<m:boxPr><m:ctrlPr><w:rPr><w:sz w:val="{0}"/><w:szCs w:val="{0}"/></w:rPr></m:ctrlPr></m:boxPr>\2\3'.format(self.正文sz*2),box格式)
        节点 = runs._element
        解析 = etree.XML(结果)[0]
        节点.append(解析)

    def 写入(self,标题xpath:str,作者xpath:str,正文xpath:str):
        if len(self.url) > 1:
            文件命名 = input('保存文件命名(不用加后缀)：')
        for a in self.url:
            html = etree.HTML(req.get(a).text)
            标题 = html.xpath(标题xpath)[0]
            作者 = html.xpath(作者xpath)[0]
            正文 = html.xpath(正文xpath)
            print('正在抓取[%s]' % 标题)
            self.写入标题和作者(标题,作者)
            self.写入正文(正文)
            if len(self.url) > 1:
                runs = self.docx.add_paragraph().add_run()
                runs.add_break(WD_BREAK.PAGE)
                地址 = os.path.join(self.储存目录dir,文件命名)+'.docx'
                self.docx.save(地址)
            else:
                地址 = os.path.join(self.储存目录dir,标题)+'.docx'
                self.docx.save(地址)
            time.sleep(1)

def main():
    # 专栏 = req.get('').text
    url = ['']
    # for a in json.loads(专栏)['data']:
    #     url.append(a['url'])
    dir = r''
    标题 = '//*[@class="Post-Main Post-NormalMain"]/header/h1/text()'
    作者 = '//*[@class="UserLink-link"]/text()'
    正文 = '//div[@options="[object Object]"]/node()'
    抓取 = 抓取文章(url,dir)
    抓取.写入(标题,作者,正文)

if __name__ == '__main__':
    main()

